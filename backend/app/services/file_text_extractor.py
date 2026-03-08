import io
import re
from pathlib import Path

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".xml",
    ".yaml",
    ".yml",
    ".log",
}
MODERN_EXCEL_EXTENSIONS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
SUPPORTED_TEXT_EXTRACTION_EXTENSIONS = (
    TEXT_EXTENSIONS | MODERN_EXCEL_EXTENSIONS | {".pdf", ".doc", ".docx"}
)


def _decode_text_bytes(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def _binary_best_effort_text(file_bytes: bytes) -> str:
    decoded = file_bytes.decode("latin-1", errors="ignore")
    chunks = re.findall(r"[A-Za-z0-9][A-Za-z0-9\s,.;:/()\-_%+]{5,}", decoded)
    return "\n".join(chunks[:5000]).strip()


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:
        raise ValueError("PDF parsing dependency missing. Install pypdf.") from exc

    reader = PdfReader(io.BytesIO(file_bytes))
    lines: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            lines.append(page_text)
    return "\n".join(lines).strip()


def _extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise ValueError("DOCX parsing dependency missing. Install python-docx.") from exc

    document = Document(io.BytesIO(file_bytes))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_values:
                lines.append(" | ".join(row_values))
    return "\n".join(lines).strip()


def _extract_text_from_xlsx(file_bytes: bytes) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore
    except ImportError as exc:
        raise ValueError("Excel parsing dependency missing. Install openpyxl.") from exc

    workbook = load_workbook(filename=io.BytesIO(file_bytes), data_only=True, read_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                lines.append("\t".join(values))
    return "\n".join(lines).strip()


def extract_text_from_file(
    file_bytes: bytes,
    filename: str | None,
    content_type: str | None = None,
) -> str:
    if not file_bytes:
        raise ValueError("Uploaded file is empty")

    suffix = Path(filename or "").suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        text = _decode_text_bytes(file_bytes).strip()
        if not text:
            raise ValueError("Uploaded text file did not contain readable text")
        return text

    if suffix == ".pdf":
        text = _extract_text_from_pdf(file_bytes)
        if not text:
            raise ValueError("PDF text extraction returned empty content")
        return text

    if suffix == ".docx":
        text = _extract_text_from_docx(file_bytes)
        if not text:
            raise ValueError("DOCX text extraction returned empty content")
        return text

    if suffix == ".doc":
        text = _binary_best_effort_text(file_bytes)
        if not text:
            raise ValueError("Legacy .doc extraction failed. Please upload .docx or .pdf.")
        return text

    if suffix in MODERN_EXCEL_EXTENSIONS:
        text = _extract_text_from_xlsx(file_bytes)
        if not text:
            raise ValueError("Excel extraction returned empty content")
        return text

    if suffix == ".xls":
        raise ValueError("Legacy .xls is not supported yet. Please convert the file to .xlsx or PDF.")

    if content_type and content_type.startswith("text/"):
        text = _decode_text_bytes(file_bytes).strip()
        if text:
            return text

    text = _binary_best_effort_text(file_bytes)
    if text:
        return text
    raise ValueError("Unsupported file format. Upload txt/csv/json/pdf/doc/docx/xlsx.")
