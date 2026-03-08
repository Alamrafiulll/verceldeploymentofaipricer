from __future__ import annotations

import csv
from decimal import Decimal
from pathlib import Path

from openpyxl import Workbook
from sqlalchemy import select

from app.db.models import Product
from app.db.session import SessionLocal


ROOT_DIR = Path(__file__).resolve().parents[3]
OUTPUT_DIR = ROOT_DIR / "docs" / "test-files"
VALID_DIR = OUTPUT_DIR / "valid"
INVALID_DIR = OUTPUT_DIR / "invalid"


def _load_product_samples(limit: int = 6) -> list[tuple[str, float]]:
    fallback = [
        ("SKU-1000", 299.0),
        ("SKU-1001", 349.0),
        ("SKU-1002", 399.0),
        ("SKU-1003", 449.0),
        ("SKU-1004", 499.0),
        ("SKU-1005", 549.0),
    ]

    try:
        db = SessionLocal()
        rows = db.execute(
            select(Product.sku, Product.list_price).order_by(Product.created_at.asc()).limit(limit)
        ).all()
        db.close()
        if not rows:
            return fallback
        samples: list[tuple[str, float]] = []
        for sku, list_price in rows:
            value = float(list_price if isinstance(list_price, Decimal) else list_price or 0)
            samples.append((sku, round(value, 2)))
        return samples
    except Exception:
        return fallback


def _escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _build_simple_pdf(lines: list[str]) -> bytes:
    y = 760
    commands = ["BT", "/F1 12 Tf", "72 760 Td"]
    for index, line in enumerate(lines):
        if index == 0:
            commands.append(f"({_escape_pdf_text(line)}) Tj")
        else:
            y -= 16
            commands.append(f"72 {y} Td")
            commands.append(f"({_escape_pdf_text(line)}) Tj")
    commands.append("ET")
    stream_text = "\n".join(commands) + "\n"
    stream_bytes = stream_text.encode("latin-1", errors="ignore")

    objects: list[bytes] = []
    objects.append(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    objects.append(b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    objects.append(
        (
            "3 0 obj\n"
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            "/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\n"
            "endobj\n"
        ).encode("ascii")
    )
    objects.append(b"4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n")
    objects.append(
        (
            f"5 0 obj\n<< /Length {len(stream_bytes)} >>\nstream\n".encode("ascii")
            + stream_bytes
            + b"endstream\nendobj\n"
        )
    )

    pdf = bytearray()
    pdf.extend(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj)

    xref_start = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_start}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(pdf)


def _write_pricebook_files(products: list[tuple[str, float]]) -> None:
    valid_rows = products[:4]
    valid_csv = VALID_DIR / "pricebook_valid_lsp.csv"
    with valid_csv.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["sku", "list_price", "notes"])
        for index, (sku, base_price) in enumerate(valid_rows):
            writer.writerow([sku, round(base_price * 0.97, 2), f"LSP test row {index + 1}"])

    valid_xlsx = VALID_DIR / "pricebook_valid_wm.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "WM Pricebook"
    ws.append(["sku", "list_price", "notes"])
    for index, (sku, base_price) in enumerate(valid_rows):
        ws.append([sku, round(base_price * 0.93, 2), f"WM test row {index + 1}"])
    wb.save(valid_xlsx)

    valid_lsp_xlsx = VALID_DIR / "pricebook_valid_lsp.xlsx"
    wb_lsp = Workbook()
    ws_lsp = wb_lsp.active
    ws_lsp.title = "LSP Pricebook"
    ws_lsp.append(["sku", "list_price", "notes"])
    for index, (sku, base_price) in enumerate(valid_rows):
        ws_lsp.append([sku, round(base_price * 0.97, 2), f"LSP test row {index + 1}"])
    wb_lsp.save(valid_lsp_xlsx)

    missing_header = INVALID_DIR / "pricebook_invalid_missing_header.csv"
    with missing_header.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["sku", "price", "notes"])
        writer.writerow([valid_rows[0][0], "120.00", "wrong header price"])

    unknown_sku = INVALID_DIR / "pricebook_invalid_unknown_sku.csv"
    with unknown_sku.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["sku", "list_price", "notes"])
        writer.writerow(["SKU-NOT-EXIST", "199.00", "should fail unknown SKU"])

    bad_price = INVALID_DIR / "pricebook_invalid_bad_price.csv"
    with bad_price.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["sku", "list_price", "notes"])
        writer.writerow([valid_rows[1][0], "ABC", "non-numeric price"])

    invalid_missing_header_xlsx = INVALID_DIR / "pricebook_invalid_missing_header.xlsx"
    wb_missing = Workbook()
    ws_missing = wb_missing.active
    ws_missing.title = "InvalidHeader"
    ws_missing.append(["sku", "price", "notes"])
    ws_missing.append([valid_rows[0][0], "120.00", "wrong header price"])
    wb_missing.save(invalid_missing_header_xlsx)

    invalid_unknown_sku_xlsx = INVALID_DIR / "pricebook_invalid_unknown_sku.xlsx"
    wb_unknown = Workbook()
    ws_unknown = wb_unknown.active
    ws_unknown.title = "UnknownSKU"
    ws_unknown.append(["sku", "list_price", "notes"])
    ws_unknown.append(["SKU-NOT-EXIST", 199.00, "should fail unknown SKU"])
    wb_unknown.save(invalid_unknown_sku_xlsx)

    invalid_bad_price_xlsx = INVALID_DIR / "pricebook_invalid_bad_price.xlsx"
    wb_bad_price = Workbook()
    ws_bad_price = wb_bad_price.active
    ws_bad_price.title = "BadPrice"
    ws_bad_price.append(["sku", "list_price", "notes"])
    ws_bad_price.append([valid_rows[1][0], "ABC", "non-numeric price"])
    wb_bad_price.save(invalid_bad_price_xlsx)


def _write_policy_files() -> None:
    memo_text = (
        "FY2025 Toiletries Bag Free Gift Campaign for DC pump water heater excluding FLUSSO series.\n"
        "Effective: 1 Jul 2025 to 16 Sep 2025.\n"
        "Not applicable for Corporate Account, Project Sales, Special Price Purchase.\n"
        "Free gift product codes: RPG-BAG-NB and RPG-BAG-GR.\n"
        "Managers may review exceptional approvals case by case.\n"
    )
    (VALID_DIR / "policy_memo_fy2025.txt").write_text(memo_text, encoding="utf-8")

    trading_terms = (
        "Trading Terms FY2025\n"
        "Dealer turnover rebate tiers apply quarterly.\n"
        "Incentive: 2% rebate above threshold A, 3% above threshold B.\n"
        "Display incentive applies when POS materials are installed.\n"
    )
    (VALID_DIR / "policy_trading_terms.doc").write_text(trading_terms, encoding="utf-8")

    try:
        from docx import Document  # type: ignore

        doc = Document()
        doc.add_heading("Campaign Memo", level=1)
        doc.add_paragraph(
            "FY2025 Toiletries Bag Free Gift Campaign for DC pump water heater excluding FLUSSO series."
        )
        doc.add_paragraph("Effective 1 Jul 2025 to 16 Sep 2025.")
        doc.add_paragraph(
            "Not applicable for Corporate Account, Project Sales, Special Price Purchase."
        )
        doc.add_paragraph("Gift SKU: RPG-BAG-NB, RPG-BAG-GR")
        doc.save(VALID_DIR / "policy_campaign.docx")
    except Exception:
        (VALID_DIR / "policy_campaign.docx").write_text(
            "Install python-docx to regenerate this file.",
            encoding="utf-8",
        )

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Campaign Rules"
    sheet.append(["field", "value"])
    sheet.append(["campaign_name", "FY2025 Toiletries Bag Free Gift Campaign"])
    sheet.append(["eligibility", "product_category=water_heater; model_type=dc_pump"])
    sheet.append(["exclusion", "series=FLUSSO"])
    sheet.append(["not_applicable_for", "corporate_account,project_sales,special_price_purchase"])
    sheet.append(["gift_skus", "RPG-BAG-NB,RPG-BAG-GR"])
    sheet.append(["effective_start", "2025-07-01"])
    sheet.append(["effective_end", "2025-09-16"])
    workbook.save(VALID_DIR / "policy_campaign_matrix.xlsx")

    pdf_lines = [
        "FY2025 Toiletries Bag Free Gift Campaign",
        "DC pump water heater eligible. FLUSSO excluded.",
        "Not applicable: Corporate Account, Project Sales, Special Price Purchase.",
        "Gift SKU: RPG-BAG-NB, RPG-BAG-GR.",
    ]
    (VALID_DIR / "policy_campaign.pdf").write_bytes(_build_simple_pdf(pdf_lines))

    (INVALID_DIR / "policy_invalid_empty.txt").write_text("", encoding="utf-8")

    invalid_policy_xlsx = INVALID_DIR / "policy_invalid_empty.xlsx"
    invalid_wb = Workbook()
    invalid_sheet = invalid_wb.active
    invalid_sheet.title = "EmptyPolicy"
    invalid_sheet.append(["field", "value"])
    invalid_wb.save(invalid_policy_xlsx)


def _write_readme(products: list[tuple[str, float]]) -> None:
    sku_preview = ", ".join(sku for sku, _ in products[:4])
    content = f"""# Test Files

These files are generated for manual upload testing in:

- `POST /api/policies/upload` (Sales/Admin)
- `POST /api/pricebooks/upload` (Sales only)

Generated SKU sample used for valid pricebooks: `{sku_preview}`.

## Valid files (`valid/`)

- `pricebook_valid_lsp.csv`: Valid CSV with `sku,list_price,notes` headers.
- `pricebook_valid_lsp.xlsx`: Valid LSP workbook for Excel-only testing.
- `pricebook_valid_wm.xlsx`: Valid XLSX with same headers.
- `policy_memo_fy2025.txt`: Memo text designed to trigger eligibility/exclusion/entitlement extraction.
- `policy_campaign.docx`: DOCX memo for policy parser test.
- `policy_trading_terms.doc`: Legacy DOC-like text for best-effort parser.
- `policy_campaign_matrix.xlsx`: Spreadsheet policy input.
- `policy_campaign.pdf`: PDF with readable text lines.

## Invalid files (`invalid/`)

- `pricebook_invalid_missing_header.csv`: Uses `price` instead of `list_price` header.
- `pricebook_invalid_missing_header.xlsx`: Same invalid header case in Excel format.
- `pricebook_invalid_unknown_sku.csv`: Contains unknown SKU.
- `pricebook_invalid_unknown_sku.xlsx`: Same unknown SKU case in Excel format.
- `pricebook_invalid_bad_price.csv`: Non-numeric `list_price`.
- `pricebook_invalid_bad_price.xlsx`: Same non-numeric list price case in Excel format.
- `policy_invalid_empty.txt`: Empty policy file.
- `policy_invalid_empty.xlsx`: Near-empty workbook to trigger invalid/weak extraction path.

## Expected behavior

- Valid files should upload successfully (if SKUs exist in database).
- Invalid files should return 400/422 with clear validation messages.
"""
    (OUTPUT_DIR / "README.md").write_text(content, encoding="utf-8")


def main() -> None:
    VALID_DIR.mkdir(parents=True, exist_ok=True)
    INVALID_DIR.mkdir(parents=True, exist_ok=True)

    products = _load_product_samples(limit=8)
    _write_pricebook_files(products)
    _write_policy_files()
    _write_readme(products)

    print(f"Generated test files in: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
